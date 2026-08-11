from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document

from app.schemas import (
    BindingStrategy,
    FileType,
    TemplateBinding,
    TemplateField,
    TemplateManifest,
)
from app.services.docx.inspect import inspect_docx
from app.services.docx.render import render_docx


def _hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _manifest() -> TemplateManifest:
    return TemplateManifest(
        template_id="fixture",
        name="Synthetic memo",
        file_type=FileType.DOCX,
        original_file="memo.docx",
        document_type="internal_memo",
        fields=[
            TemplateField(
                id=field_id,
                label=field_id,
                binding=TemplateBinding(
                    strategy=BindingStrategy.PLACEHOLDER,
                    placeholder=f"{{{{{field_id}}}}}",
                ),
            )
            for field_id in ("subject", "recipient", "body", "signer_name", "signer_title")
        ],
    )


def test_render_replaces_split_placeholders_without_mutating_original(
    docx_fixture: Path, tmp_path: Path
) -> None:
    before_hash = _hash(docx_fixture)
    before = inspect_docx(docx_fixture)
    output = tmp_path / "rendered.docx"

    warnings = render_docx(
        docx_fixture,
        output,
        _manifest(),
        {
            "subject": "ขออนุมัติจัดซื้อเครื่องสำรองไฟฟ้า",
            "recipient": "ผู้อำนวยการ",
            "body": "ข้อความภาษาไทยยังคงสมบูรณ์",
            "signer_name": "ผู้ทดสอบระบบ",
            "signer_title": "เภสัชกร",
        },
    )

    assert warnings == []
    assert _hash(docx_fixture) == before_hash
    assert output.is_file()
    rendered = Document(output)
    all_text = "\n".join(
        [paragraph.text for paragraph in rendered.paragraphs]
        + [cell.text for table in rendered.tables for row in table.rows for cell in row.cells]
    )
    assert "ข้อความภาษาไทยยังคงสมบูรณ์" in all_text
    assert "{{body}}" not in all_text

    after = inspect_docx(output)
    assert after["sections"] == before["sections"]
    assert (
        after["headers"][0]["paragraphs"][0]["text"]
        == before["headers"][0]["paragraphs"][0]["text"]
    )
    assert (
        after["footers"][0]["paragraphs"][0]["text"]
        == before["footers"][0]["paragraphs"][0]["text"]
    )
    assert after["package"]["preservation_checksums"] == before["package"]["preservation_checksums"]
