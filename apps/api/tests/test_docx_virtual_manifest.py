from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


def _create_thai_candidate_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("เรื่อง: ขออนุมัติโครงการ")
    document.add_paragraph("เรียน ผู้อำนวยการ")
    document.add_paragraph("วันที่ 20 สิงหาคม 2569")
    document.add_paragraph("ที่ กค 0001/2569")
    document.save(path)
    return path


def _upload(client: TestClient, path: Path, filename: str) -> str:
    with path.open("rb") as handle:
        response = client.post(
            "/api/templates/upload",
            files={"file": (filename, handle, "application/octet-stream")},
        )
    assert response.status_code == 201, response.text
    return response.json()["id"]


def _analyze(client: TestClient, template_id: str) -> dict:
    response = client.post(f"/api/templates/{template_id}/analyze")
    assert response.status_code == 200, response.text
    return response.json()


def _native_manifest(template_id: str, *, block_id: str = "body.p0", anchor: str = "เรื่อง") -> dict:
    return {
        "template_id": template_id,
        "name": "Reviewed Thai memo",
        "file_type": "docx",
        "original_file": "ignored-by-server.docx",
        "document_type": "internal_memo",
        "fields": [
            {
                "id": "subject",
                "label": "เรื่อง",
                "required": False,
                "content_type": "text",
                "binding": {
                    "strategy": "docx_block",
                    "block_id": block_id,
                    "anchor": anchor,
                    "mode": "after_anchor",
                },
            }
        ],
    }


def test_docx_analysis_exposes_proposals_without_accepting_them(tmp_path: Path) -> None:
    source = _create_thai_candidate_docx(tmp_path / "thai-memo.docx")
    with TestClient(app) as client:
        template_id = _upload(client, source, "thai-memo.docx")
        analyzed = _analyze(client, template_id)

    assert analyzed["manifest"]["fields"] == []
    assert [candidate["field_id"] for candidate in analyzed["analysis"]["candidates"]] == [
        "subject",
        "recipient",
        "date",
        "document_number",
    ]
    assert analyzed["analysis"]["candidates"][0]["source_text"] == "เรื่อง: ขออนุมัติโครงการ"


def test_valid_reviewed_docx_binding_is_persisted_and_native_render_works(
    tmp_path: Path,
) -> None:
    source = _create_thai_candidate_docx(tmp_path / "reviewed.docx")
    before_hash = hashlib.sha256(source.read_bytes()).hexdigest()
    with TestClient(app) as client:
        template_id = _upload(client, source, "reviewed.docx")
        _analyze(client, template_id)
        saved = client.put(
            f"/api/templates/{template_id}/manifest",
            json=_native_manifest(template_id),
        )
        render = client.post(
            "/api/generations/render",
            json={"template_id": template_id, "content": {"subject": "ใหม่"}},
        )

    assert saved.status_code == 200, saved.text
    assert saved.json()["manifest"]["fields"][0]["binding"]["strategy"] == "docx_block"
    assert render.status_code == 200, render.text
    output = client.get(render.json()["download_url"])
    assert output.status_code == 200
    rendered = Document(BytesIO(output.content))
    assert rendered.paragraphs[0].text == "เรื่อง: ใหม่"
    assert hashlib.sha256(source.read_bytes()).hexdigest() == before_hash


def test_unknown_block_id_is_rejected(tmp_path: Path) -> None:
    source = _create_thai_candidate_docx(tmp_path / "unknown.docx")
    with TestClient(app) as client:
        template_id = _upload(client, source, "unknown.docx")
        _analyze(client, template_id)
        response = client.put(
            f"/api/templates/{template_id}/manifest",
            json=_native_manifest(template_id, block_id="body.p99"),
        )

    assert response.status_code == 422
    assert "Unknown DOCX block_id" in response.json()["detail"]


def test_invalid_anchor_is_rejected(tmp_path: Path) -> None:
    source = _create_thai_candidate_docx(tmp_path / "anchor.docx")
    with TestClient(app) as client:
        template_id = _upload(client, source, "anchor.docx")
        _analyze(client, template_id)
        response = client.put(
            f"/api/templates/{template_id}/manifest",
            json=_native_manifest(template_id, anchor="เรียน"),
        )

    assert response.status_code == 422
    assert "is not valid" in response.json()["detail"]


def test_wrong_strategy_is_rejected_for_each_file_type(
    tmp_path: Path, xlsx_fixture: Path
) -> None:
    docx_source = _create_thai_candidate_docx(tmp_path / "wrong-docx.docx")
    with TestClient(app) as client:
        docx_id = _upload(client, docx_source, "wrong-docx.docx")
        _analyze(client, docx_id)
        docx_response = client.put(
            f"/api/templates/{docx_id}/manifest",
            json={
                **_native_manifest(docx_id),
                "fields": [
                    {
                        "id": "subject",
                        "label": "Subject",
                        "required": False,
                        "content_type": "text",
                        "binding": {"strategy": "cell", "cell": "A1"},
                    }
                ],
            },
        )
        xlsx_id = _upload(client, xlsx_fixture, "comparison.xlsx")
        xlsx_response = client.put(
            f"/api/templates/{xlsx_id}/manifest",
            json={
                "template_id": xlsx_id,
                "name": "Wrong XLSX binding",
                "file_type": "xlsx",
                "original_file": "comparison.xlsx",
                "fields": [
                    {
                        "id": "subject",
                        "label": "Subject",
                        "required": False,
                        "content_type": "text",
                        "binding": {
                            "strategy": "docx_block",
                            "block_id": "body.p0",
                            "anchor": "เรื่อง",
                            "mode": "after_anchor",
                        },
                    }
                ],
            },
        )

    assert docx_response.status_code == 422
    assert "support placeholder and docx_block" in docx_response.json()["detail"]
    assert xlsx_response.status_code == 422
    assert "support cell bindings only" in xlsx_response.json()["detail"]


def test_template_manifests_are_isolated(tmp_path: Path) -> None:
    first_source = _create_thai_candidate_docx(tmp_path / "first.docx")
    second_source = _create_thai_candidate_docx(tmp_path / "second.docx")
    with TestClient(app) as client:
        first_id = _upload(client, first_source, "first.docx")
        second_id = _upload(client, second_source, "second.docx")
        _analyze(client, first_id)
        _analyze(client, second_id)
        first_saved = client.put(
            f"/api/templates/{first_id}/manifest",
            json=_native_manifest(first_id),
        )
        second = client.get(f"/api/templates/{second_id}")

    assert first_saved.status_code == 200
    assert second.status_code == 200
    assert second.json()["manifest"]["fields"] == []
