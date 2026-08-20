from __future__ import annotations

import hashlib
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas import FileType, TemplateBinding, TemplateField, TemplateManifest
from app.services.docx.inspect import inspect_docx
from app.services.docx.render import DocxRenderError, render_docx


def _manifest(*fields: TemplateField) -> TemplateManifest:
    return TemplateManifest(
        template_id="native-fixture",
        name="Native fixture",
        file_type=FileType.DOCX,
        original_file="native.docx",
        fields=list(fields),
    )


def _native_field(
    field_id: str,
    block_id: str,
    anchor: str,
    *,
    required: bool = True,
    mode: str = "after_anchor",
) -> TemplateField:
    return TemplateField(
        id=field_id,
        label=field_id,
        required=required,
        binding=TemplateBinding(
            strategy="docx_block",
            block_id=block_id,
            anchor=anchor,
            mode=mode,
        ),
    )


def _create_native_docx(path: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "เรียน: ส่วนหัวเดิม"
    section.footer.paragraphs[0].text = "ที่\tท้ายเอกสารเดิม"

    paragraph = document.add_paragraph()
    paragraph.style = "Title"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor = paragraph.add_run("เรื่อง")
    anchor.bold = True
    paragraph.add_run(": ")
    paragraph.add_run("ขออนุมัติเดิม")

    split_anchor = document.add_paragraph()
    split_anchor.add_run("เร")
    split_anchor.add_run("ียน")
    split_anchor.add_run(" ")
    split_anchor.add_run("ผู้อำนวย")
    split_anchor.add_run("การเดิม")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "วันที่\t20 สิงหาคม 2569"
    nested = table.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "เรื่อง: nested เดิม"
    document.save(path)
    return path


def _render(source: Path, output: Path, *fields: TemplateField, content: dict) -> None:
    render_docx(source, output, _manifest(*fields), content)


def test_after_anchor_replaces_body_value_and_preserves_formatting(tmp_path: Path) -> None:
    source = _create_native_docx(tmp_path / "body.docx")
    output = tmp_path / "rendered.docx"
    before = inspect_docx(source)
    _render(
        source,
        output,
        _native_field("subject", "body.p0", "เรื่อง"),
        content={"subject": "ขออนุมัติใหม่"},
    )

    rendered = Document(output)
    paragraph = rendered.paragraphs[0]
    assert paragraph.text == "เรื่อง: ขออนุมัติใหม่"
    assert paragraph.style.name == before["paragraphs"][0]["style"]
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph.runs[0].text == "เรื่อง"
    assert paragraph.runs[0].bold is True
    assert len(paragraph.runs) >= 3
    assert paragraph.runs[2].text == "ขออนุมัติใหม่"
    assert paragraph.runs[2].bold is not True


def test_after_anchor_handles_split_anchor_and_split_value(tmp_path: Path) -> None:
    source = _create_native_docx(tmp_path / "split.docx")
    output = tmp_path / "rendered.docx"
    _render(
        source,
        output,
        _native_field("recipient", "body.p1", "เรียน"),
        content={"recipient": "ผู้บริหารใหม่"},
    )

    assert Document(output).paragraphs[1].text == "เรียน ผู้บริหารใหม่"
    assert len(Document(output).paragraphs[1].runs) >= 3


@pytest.mark.parametrize(
    ("block_id", "anchor", "value", "expected"),
    [
        ("table.0.row.0.cell.0.p0", "วันที่", "วันที่ใหม่", "วันที่\tวันที่ใหม่"),
        ("header.0.p0", "เรียน", "หัวข้อใหม่", "เรียน: หัวข้อใหม่"),
        ("footer.0.p0", "ที่", "ท้ายใหม่", "ที่\tท้ายใหม่"),
    ],
)
def test_after_anchor_resolves_table_header_and_footer_blocks(
    tmp_path: Path,
    block_id: str,
    anchor: str,
    value: str,
    expected: str,
) -> None:
    source = _create_native_docx(tmp_path / f"{anchor}.docx")
    output = tmp_path / f"{anchor}-rendered.docx"
    _render(
        source,
        output,
        _native_field("field", block_id, anchor),
        content={"field": value},
    )

    document = Document(output)
    if block_id.startswith("table"):
        actual = document.tables[0].cell(0, 0).paragraphs[0].text
    elif block_id.startswith("header"):
        actual = document.sections[0].header.paragraphs[0].text
    else:
        actual = document.sections[0].footer.paragraphs[0].text
    assert actual == expected


def test_missing_value_clears_stale_value_and_warns_when_required(tmp_path: Path) -> None:
    source = _create_native_docx(tmp_path / "missing.docx")
    output = tmp_path / "rendered.docx"
    manifest = _manifest(_native_field("subject", "body.p0", "เรื่อง"))

    warnings = render_docx(source, output, manifest, {})

    assert warnings == ["Required field 'subject' had no value."]
    assert Document(output).paragraphs[0].text == "เรื่อง: "


def test_stale_block_anchor_and_mode_fail_clearly(tmp_path: Path) -> None:
    source = _create_native_docx(tmp_path / "invalid.docx")
    output = tmp_path / "rendered.docx"

    with pytest.raises(DocxRenderError, match="could not be resolved"):
        _render(
            source,
            output,
            _native_field("subject", "body.p99", "เรื่อง"),
            content={"subject": "ใหม่"},
        )
    with pytest.raises(DocxRenderError, match="was not found"):
        _render(
            source,
            output,
            _native_field("subject", "body.p0", "เรียน"),
            content={"subject": "ใหม่"},
        )
    with pytest.raises(ValueError, match="after_anchor"):
        _native_field("subject", "body.p0", "เรื่อง", mode="unsupported")


def test_render_copies_source_and_preserves_package_parts(tmp_path: Path) -> None:
    source = _create_native_docx(tmp_path / "preserve.docx")
    output = tmp_path / "rendered.docx"
    before_hash = hashlib.sha256(source.read_bytes()).hexdigest()
    before = inspect_docx(source)

    _render(
        source,
        output,
        _native_field("subject", "body.p0", "เรื่อง"),
        content={"subject": "ใหม่"},
    )
    after = inspect_docx(output)

    assert hashlib.sha256(source.read_bytes()).hexdigest() == before_hash
    assert after["sections"] == before["sections"]
    assert after["package"]["preservation_checksums"] == before["package"]["preservation_checksums"]
    assert after["package"]["parts"] == before["package"]["parts"]


def test_placeholder_and_native_bindings_render_together(tmp_path: Path) -> None:
    source = _create_native_docx(tmp_path / "mixed.docx")
    paragraph = Document(source).add_paragraph()
    paragraph.add_run("{{footer_note}}")
    document = Document(source)
    document.add_paragraph("{{footer_note}}")
    document.save(source)
    output = tmp_path / "rendered.docx"
    _render(
        source,
        output,
        _native_field("subject", "body.p0", "เรื่อง"),
        TemplateField(
            id="footer_note",
            label="footer_note",
            binding=TemplateBinding(strategy="placeholder", placeholder="{{footer_note}}"),
        ),
        content={"subject": "ใหม่", "footer_note": "หมายเหตุ"},
    )

    assert "เรื่อง: ใหม่" in Document(output).paragraphs[0].text
    assert "หมายเหตุ" in "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
