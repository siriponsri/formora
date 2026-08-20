from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.docx.candidates import detect_candidates
from app.services.docx.inspect import inspect_docx


def _create_structural_docx(path: Path, *, with_anchors: bool = True) -> Path:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "ส่วนหัวเอกสาร"
    section.footer.paragraphs[0].text = "ส่วนท้ายเอกสาร"

    document.add_paragraph("บรรทัดแรก")
    document.add_paragraph(
        "เรื่อง: ขออนุมัติโครงการ" if with_anchors else "รายละเอียดโครงการ"
    )
    document.add_paragraph("เรียน ผู้อำนวยการ" if with_anchors else "เรียนรู้จากเอกสาร")
    document.add_paragraph("วันที่ 20 สิงหาคม 2569" if with_anchors else "วันที่จัดทำ")
    document.add_paragraph("ที่ กค 0001/2569" if with_anchors else "ที่ประชุมมีมติ")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "ตาราง"
    table.cell(0, 1).text = "เรื่อง: หัวข้อในตาราง" if with_anchors else "ข้อมูลทั่วไป"
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(path)
    return path


def _all_paragraph_maps(document_map: dict) -> list[dict]:
    paragraphs = list(document_map["paragraphs"])
    for table in document_map["tables"]:
        paragraphs.extend(
            paragraph
            for row in table["rows"]
            for cell in row["cells"]
            for paragraph in cell["paragraphs"]
        )
    for story_name in ("headers", "footers"):
        for story in document_map[story_name]:
            paragraphs.extend(story["paragraphs"])
    return paragraphs


def test_inspection_assigns_stable_ids_to_body_table_header_and_footer(
    tmp_path: Path,
) -> None:
    document_map = inspect_docx(_create_structural_docx(tmp_path / "structure.docx"))

    assert [paragraph["block_id"] for paragraph in document_map["paragraphs"]] == [
        "body.p0",
        "body.p1",
        "body.p2",
        "body.p3",
        "body.p4",
    ]
    assert document_map["tables"][0]["rows"][0]["cells"][1]["paragraphs"][0][
        "block_id"
    ] == "table.0.row.0.cell.1.p0"
    assert document_map["headers"][0]["paragraphs"][0]["block_id"] == "header.0.p0"
    assert document_map["footers"][0]["paragraphs"][0]["block_id"] == "footer.0.p0"


def test_block_ids_are_unique_and_inspection_is_deterministic(tmp_path: Path) -> None:
    path = _create_structural_docx(tmp_path / "stable.docx")
    first = inspect_docx(path)
    second = inspect_docx(path)
    first_ids = [paragraph["block_id"] for paragraph in _all_paragraph_maps(first)]
    second_ids = [paragraph["block_id"] for paragraph in _all_paragraph_maps(second)]

    assert len(first_ids) == len(set(first_ids))
    assert first_ids == second_ids
    assert first == second


def test_detects_explicit_thai_anchor_candidates(tmp_path: Path) -> None:
    document_map = inspect_docx(_create_structural_docx(tmp_path / "anchors.docx"))

    candidates = detect_candidates(document_map)

    assert [(candidate["field_id"], candidate["block_id"]) for candidate in candidates] == [
        ("subject", "body.p1"),
        ("recipient", "body.p2"),
        ("date", "body.p3"),
        ("document_number", "body.p4"),
        ("subject", "table.0.row.0.cell.1.p0"),
    ]
    assert all(candidate["label"] == candidate["anchor"] for candidate in candidates)
    assert all(candidate["confidence"] == 0.99 for candidate in candidates)


def test_returns_no_candidates_without_clear_anchor_evidence(tmp_path: Path) -> None:
    document_map = inspect_docx(
        _create_structural_docx(tmp_path / "no-anchors.docx", with_anchors=False)
    )

    assert detect_candidates(document_map) == []


def test_candidate_references_are_real_inspected_blocks(tmp_path: Path) -> None:
    document_map = inspect_docx(_create_structural_docx(tmp_path / "references.docx"))
    block_ids = {paragraph["block_id"] for paragraph in _all_paragraph_maps(document_map)}

    assert all(candidate["block_id"] in block_ids for candidate in detect_candidates(document_map))
