from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.page import PageMargins


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def create_sample_docx(path: str | Path) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.0)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = document.styles["Normal"]
    normal.font.name = "TH Sarabun New"
    normal.font.size = Pt(16)

    header = section.header.paragraphs[0]
    header.text = "FORMORA · SYNTHETIC TEMPLATE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer.paragraphs[0]
    footer.text = "เอกสารตัวอย่างสำหรับการทดสอบเท่านั้น"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("บันทึกข้อความ")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(49, 46, 129)

    metadata = document.add_table(rows=2, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata.style = "Table Grid"
    metadata.columns[0].width = Cm(3.5)
    metadata.columns[1].width = Cm(13.5)
    metadata.cell(0, 0).text = "เรื่อง"
    metadata.cell(0, 1).text = "{{subject}}"
    metadata.cell(1, 0).text = "เรียน"
    metadata.cell(1, 1).text = "{{recipient}}"
    for row in metadata.rows:
        _set_cell_shading(row.cells[0], "EEF2FF")
        row.cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()
    body = document.add_paragraph()
    body.paragraph_format.first_line_indent = Cm(1.25)
    body.paragraph_format.line_spacing = 1.15
    # Split the placeholder across runs intentionally to exercise the renderer.
    body.add_run("{{bo").font.name = "TH Sarabun New"
    body.add_run("dy}}").font.name = "TH Sarabun New"

    document.add_paragraph()
    signature = document.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature.add_run("ลงชื่อ ....................................................\n")
    signature.add_run("({{signer_name}})\n").bold = True
    signature.add_run("{{signer_title}}")

    document.core_properties.title = "Synthetic internal memo template"
    document.core_properties.subject = "Formora preservation fixture"
    document.save(target)
    return target


def create_sample_xlsx(path: str | Path) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Price Comparison"
    worksheet.merge_cells("A1:F2")
    worksheet["A1"] = "ตารางเปรียบเทียบราคา · เอกสารสังเคราะห์"
    worksheet["A1"].font = Font(name="Aptos", size=18, bold=True, color="FFFFFF")
    worksheet["A1"].fill = PatternFill("solid", fgColor="4F46E5")
    worksheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
    worksheet.row_dimensions[1].height = 30
    worksheet.row_dimensions[2].height = 16

    worksheet["A3"] = "ชื่อโครงการ"
    worksheet["B3"] = "โครงการตัวอย่าง"
    worksheet.merge_cells("B3:F3")
    headers = ["รายการ", "จำนวน", "หน่วย", "ผู้เสนอ A", "ผู้เสนอ B", "ผู้เสนอ C"]
    thin = Side(style="thin", color="CBD5E1")
    for column, value in enumerate(headers, start=1):
        cell = worksheet.cell(row=5, column=column, value=value)
        cell.font = Font(bold=True, color="1E293B")
        cell.fill = PatternFill("solid", fgColor="E0E7FF")
        cell.alignment = Alignment(horizontal="center")
        cell.border = Border(top=thin, right=thin, bottom=thin, left=thin)

    worksheet["A6"] = "เครื่องสำรองไฟฟ้า 3 kVA"
    worksheet["B6"] = 1
    worksheet["C6"] = "เครื่อง"
    worksheet["D6"] = 45000
    worksheet["E6"] = 46500
    worksheet["F6"] = 47200
    worksheet["A8"] = "ราคาต่ำสุด"
    worksheet["D8"] = "=MIN(D6:F6)"
    worksheet["A10"] = "ข้อมูลสำหรับการ binding"
    worksheet["D10"] = 45000
    worksheet["E10"] = 46500
    worksheet["F10"] = 47200
    for row in worksheet.iter_rows(min_row=6, max_row=10, min_col=1, max_col=6):
        for cell in row:
            cell.border = Border(top=thin, right=thin, bottom=thin, left=thin)

    widths = {"A": 34, "B": 12, "C": 14, "D": 18, "E": 18, "F": 18}
    for column, width in widths.items():
        worksheet.column_dimensions[column].width = width
    worksheet.row_dimensions[6].height = 26
    worksheet.freeze_panes = "A5"
    worksheet.print_area = "A1:F10"
    worksheet.page_setup.orientation = "landscape"
    worksheet.page_setup.fitToWidth = 1
    worksheet.page_margins = PageMargins(left=0.3, right=0.3, top=0.5, bottom=0.5)
    worksheet.oddFooter.center.text = "Synthetic fixture · Formora"
    workbook.save(target)
    return target
